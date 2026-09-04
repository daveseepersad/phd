# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx"]
# ///
"""Generate the formal advisor-review Word document for the dissertation
problem-statement literature review. Content mirrors thesis.md (verified
quotes, pages, and citations from the 2026-08-30 saturated review run)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "Seepersad-Problem-Statement-Literature-Review.docx"

# --------------------------------------------------------------------------
# Content
# --------------------------------------------------------------------------

TITLE = ("Dissertation Topic Distillation:\nProblem Statements from a "
         "Systematic Literature Review of\nMulti-Agent LLM Software Development")
SUBTITLE = "Submitted for Advisor Review"
AUTHOR = "Dave Seepersad"
PROGRAM = "Ph.D. in Computer Science\nNova Southeastern University"
DATE = "September 1, 2026"

PURPOSE_PARAS = [
    "Large language model (LLM) based multi-agent systems have become a dominant "
    "architectural pattern in automated software engineering. In these systems, "
    "role-specialized agents — planner, coder, tester, reviewer — collaborate "
    "through natural-language handoffs to perform requirements analysis, code "
    "generation, testing, and repair, on the premise that decomposition and "
    "specialization improve outcomes over a single generalist agent. The premise is "
    "increasingly load-bearing in both research and industrial tooling, yet the "
    "empirical record behind it is contradictory: controlled studies report "
    "multi-agent configurations that beat, tie, and substantially lose to matched "
    "single-agent baselines, often within the same experiment when the base model "
    "or task family changes.",

    "This document distills a completed, saturated systematic literature review of "
    "that evidence into the three strongest research-worthy problems the review "
    "surfaced. Its purpose is to support topic selection for the Dissertation Idea "
    "Paper: each candidate topic is presented with (1) a literature-supported "
    "problem statement; (2) verbatim, page-anchored supporting quotations from the "
    "reviewed papers; (3) a proposed controlled experiment that could resolve or "
    "materially advance the problem; and (4) published evidence that the proposed "
    "approach is viable at dissertation scale. Each topic begins on a new page and "
    "closes with its own reference list so the topics can be reviewed independently.",

    "The review corpus was assembled from the ACM Digital Library, IEEE Xplore, "
    "OpenAlex, Crossref, and Google Scholar, with backward and forward citation "
    "snowballing from core papers. In the final run, 522 records were screened at "
    "the abstract level, 207 were selected as relevant, and 48 papers were read in "
    "full text. Reading stopped under a preregistered concept-saturation rule: at "
    "least 20 full texts read, no unread core papers, and five consecutive papers "
    "contributing no new evidence domain across a fixed 20-domain taxonomy. Every "
    "quotation reproduced in this document was mechanically verified as an exact "
    "passage of the source paper's extracted text, and every page number and "
    "citation was checked against the paper itself. Quotations are lightly "
    "normalized for PDF extraction artifacts (line-break hyphenation, dropped "
    "spaces); wording is unchanged. Sources not yet peer reviewed are flagged "
    "[preprint]; each topic's problem statement is anchored by refereed venues, "
    "with preprints in supporting roles.",

    "A note on method: the review workflow — its stages, screening criteria, "
    "citation-chaining strategy, stopping rule, and verification requirements — is "
    "my own design, executed with AI automation for the high-volume mechanical "
    "steps (multi-database search, retrieval, text extraction, abstract triage, "
    "structured note-taking, and quotation cross-checking). All automated judgments "
    "are recorded in auditable artifacts in the project repository. This document "
    "is a working distillation of those artifacts; the dissertation documents "
    "themselves will be authored by me, and the use of automation will be disclosed "
    "under NSU's Certification of Authorship.",
]

TOPICS = [
    # ------------------------------------------------------------------ T1
    {
        "heading": "Topic 1 — Resource-Bounded Adaptive Delegation",
        "proposed_title": ("Resource-Bounded Adaptive Delegation: Determining When "
                           "Specialist Multi-Agent Delegation Earns Its Inference Cost "
                           "in LLM-Based Software Engineering"),
        "problem": [
            "LLM-based multi-agent systems for software engineering delegate work to "
            "role-specialized agents on the premise that decomposition improves "
            "outcomes, yet the empirical record is contradictory: a two-agent "
            "Developer–Tester workflow beats a single agent with one model but loses "
            "with another, and adding a Designer role degrades performance sharply "
            "(Zeng et al., 2025); Waterfall-style role pipelines cut class-level "
            "pass@1 by roughly 38–40% for two of three backends while helping the "
            "third (Shafin et al., 2026); a single agent matched or numerically "
            "exceeded a four-specialist orchestration's decision accuracy "
            "(differences not statistically significant) while using roughly "
            "one-fifth the latency and tokens (Sanabria, 2026); and a single-agent "
            "screening configuration outperformed every multi-agent alternative "
            "(Radeva et al., 2026).",

            "Critically, these mixed results cannot be adjudicated because existing "
            "comparisons confound architecture with resources and implementation: "
            "Zeng et al. (2025) name “confounding implementation variables” as the "
            "central obstacle to fair architecture comparison; Agha and Miqdad "
            "(2026) concede their observed gaps “may reflect how sensitive each "
            "pattern is to prompt quality rather than architectural design alone” "
            "and report architectures differing 1.2×–4.6× in token spend; Li (2026) "
            "runs the corpus's tightest same-model, same-seed ablation yet still "
            "gives the full framework three times the single agent's tokens; and "
            "Orogat, Rostam, and Mansour (2026) show orchestration overhead alone "
            "— independent of task or model — is the dominant scalability "
            "constraint.",

            "The problem has therefore evolved from “are multi-agent systems "
            "better?” to a sharper, unanswered question: under a fixed inference "
            "budget, when does spawning a specialist earn back its communication "
            "and coordination cost? No study in this corpus (i) equalizes "
            "token/call budgets across single-agent and multi-agent arms so that "
            "architecture is the only manipulated variable, or (ii) tests a runtime "
            "policy that decides per task, per state, whether delegation's expected "
            "marginal benefit exceeds its marginal cost. The field's own authors "
            "point at the missing capability — “build adaptive pipelines” (Shafin "
            "et al., 2026), informed routing “left for future work” (Huang et al., "
            "2026), diminishing returns implying an optimal allocation point (Yang, "
            "Chai, & Zhang, 2025) — but none has built or falsifiably tested one.",
        ],
        "evidence": [
            ("Fair architecture comparison is blocked by confounds; the most "
             "realistic end-to-end development benchmark states it directly.",
             "the scientific evaluation of these systems is hampered by significant "
             "challenges, including overly simplistic benchmarks and the difficulty "
             "of conducting fair comparisons between different agent architectures "
             "due to confounding implementation variables.",
             "Zeng, Li, Xie, Ye, & Zhang (2025), p. 1 [preprint]"),
            ("Decomposition can be actively harmful, and no static choice is "
             "universally right.",
             "A single-agent approach is not necessarily optimal. A well-structured "
             "division of tasks and agents can effectively reduce task complexity, "
             "whereas an inappropriate decomposition can substantially degrade agent "
             "performance.",
             "Zeng et al. (2025), p. 8 [preprint]"),
            ("Delegation's cost can buy nothing — peer-reviewed, 100 scenarios, "
             "paired statistics with Holm–Bonferroni correction.",
             "The most reliable difference was computational efficiency: the "
             "single-agent architecture required substantially lower latency and "
             "token usage than the prompt-defined multi-agent orchestration "
             "architecture.",
             "Sanabria (2026), p. 1, Frontiers in Robotics and AI"),
            ("The payoff is model-dependent and unpredictable from architecture "
             "alone — published ablation on 100 ClassEval tasks.",
             "multi-agent LLM workflows do not uniformly enhance class-level code "
             "correctness … pass@1 accuracy declines for GPT-4o-Mini (−37.8%) and "
             "DeepSeek-Chat (−39.8%), while Claude-3.5-Haiku improves (+9.5%), "
             "indicating model-dependent [effects]",
             "Shafin, Rafi, Li, & Chen (2026), p. 2, PROMISE '26"),
            ("Added agents do not reliably add value.",
             "Model selection was the primary determinant of screening performance, "
             "outweighing strategy selection. The single-agent strategy with Qwen "
             "2.5 7B in few-shot mode achieved the best overall performance … "
             "outperforming all multi-agent alternatives.",
             "Radeva, Noncheva, Doukovska, & Popchev (2026), p. 24 [preprint]"),
            ("Per-token value of delegation varies by an order of magnitude — 270 "
             "matched executions.",
             "The cost-quality relationship was strongly asymmetric: sequential's "
             "1.2× token cost relative to single-agent produced a 11.1 "
             "percentage-point pass-rate gain, while hierarchical's 4.6× token cost "
             "produced a 26.7 percentage-point pass-rate loss.",
             "Agha & Miqdad (2026), p. 81"),
            ("Existing comparisons admit they do not isolate architecture.",
             "the performance gaps we observed may reflect how sensitive each "
             "pattern is to prompt quality rather than architectural design alone.",
             "Agha & Miqdad (2026), p. 77"),
            ("Every unnecessary delegation has a measurable price paid by default.",
             "Principle 1: Orchestration overhead is the dominant scalability "
             "constraint. Keep orchestration shallow; add coordination only when "
             "strictly required.",
             "Orogat, Rostam, & Mansour (2026), p. 10 [preprint]"),
        ],
        "experiment": {
            "Design": (
                "A within-task, budget-matched factorial experiment testing whether a "
                "resource-bounded adaptive delegation policy Pareto-dominates static "
                "architectures. H1: at each fixed total-token budget B, an adaptive "
                "controller that delegates to a specialist only when a calibrated "
                "estimate of marginal benefit exceeds marginal cost achieves higher "
                "task success than every static arm and a random-delegation control."),
            "Conditions": [
                "C1: static single agent (full context, full budget B)",
                "C2: static sequential pipeline (Planner→Coder→Tester, shared budget B)",
                "C3: static hierarchical/supervisor topology (shared budget B)",
                "C4: adaptive delegation — starts single-agent; the controller spawns "
                "a specialist mid-task only when estimated marginal benefit exceeds "
                "marginal cost given remaining budget",
                "C5: random-delegation control at C4's delegation frequency — "
                "isolates the value of adaptivity from the value of occasional "
                "delegation",
                "Each condition crossed with three total-token budget tiers (e.g., "
                "16k/64k/256k) and 3–5 fixed seeds",
            ],
            "Tasks and data": (
                "SWE-bench Lite (repository bug-fixing; stratified 100-task "
                "factorial), ClassEval (100 class-level tasks, matching Shafin et "
                "al., 2026), and TestEval/QuixBugs (matching Agha & Miqdad, 2026, "
                "and Li, 2026), plus a temporally held-out post-cutoff slice (Zeng "
                "et al.'s 2025-Q1 partition method) to guard against benchmark "
                "contamination."),
            "Controls": (
                "One pinned model version, fixed decoding parameters, one tool "
                "harness and sandbox, and a hard metered token ceiling across all "
                "agents — removing the exact confound named by Zeng et al., Agha & "
                "Miqdad, and Li. Prompts share an instruction core with equal-effort "
                "role additions published verbatim, plus a prompt-swap ablation. "
                "Architecture (and the policy) is the only manipulated variable."),
            "Metrics": [
                "Task success at fixed budget per tier; success per 1,000 tokens",
                "Area under the budget–performance curve",
                "Tokens, API calls, wall-clock latency, monetary cost (all-agent inclusive)",
                "Delegation precision/recall against counterfactual oracle labels "
                "(re-running delegated states without delegation and vice versa)",
                "Cross-run consistency across seeds; paired McNemar/Wilcoxon tests "
                "with Holm correction and bootstrap effect sizes",
            ],
            "Feasibility": (
                "No LLM training — the controller is a small logistic/heuristic "
                "model over trace features. Corpus cost anchors ($0.0056–$7.05 per "
                "task or project; Li, 2026; Zeng et al., 2025; Agha & Miqdad, 2026) "
                "put the full factorial at roughly $3–6k of API spend, with hard "
                "per-task ceilings bounding worst-case cost. Estimated timeline "
                "10–13 months (instrumentation and pilot; calibration and main "
                "factorial; counterfactual and statistical analysis)."),
        },
        "viability": [
            ("The controlled methodology is proven practical at single-student "
             "scale; the remaining flaw (unmatched token budgets) is exactly what "
             "the proposed budget metering fixes.",
             "Four conditions were therefore run on the same forty programs with "
             "the same model, the same five seeds, and the same iteration bound.",
             "Li (2026), p. 17, Transactions on Computing Science"),
            ("The core mechanism is an acknowledged, unclaimed gap.",
             "more informed routing strategies based on model specialization or "
             "historical performance may further accelerate convergence and improve "
             "computational efficiency, which we leave for future work.",
             "Huang, Ye, Sun, Feng, & Liu (2026), p. 8, ACM TOSEM"),
            ("Direct corpus endorsement of the dissertation's central artifact.",
             "For researchers: build adaptive pipelines, add cross-stage "
             "consistency checks, and prioritize benchmark specifications when "
             "stages conflict.",
             "Shafin et al. (2026), p. 9, PROMISE '26"),
            ("An interior optimum exists to find.",
             "However, the improvements are incremental, with diminishing returns "
             "as the system scales up.",
             "Yang, Chai, & Zhang (2025), p. 23 [preprint]"),
            ("Measurable upside, and the field's own framing of the open question.",
             "a 'best-of-N' strategy is a feasible path to improving overall "
             "performance … given the high computational cost of a full agent run, "
             "the cost-effectiveness of such an approach becomes a critical concern "
             "that requires careful optimization.",
             "Zeng et al. (2025), p. 8 [preprint]"),
        ],
        "assessment": {
            "Novelty": (
                "No corpus paper performs a budget-equalized single-versus-multi "
                "comparison in software engineering, and none tests a runtime "
                "cost-conditioned delegation policy. Both contributions are "
                "individually absent from the corpus."),
            "Falsifiability": (
                "H1 is rejected if the static single agent matches the adaptive arm "
                "at every budget tier (plausible — single agents won in Sanabria, "
                "2026; Radeva et al., 2026; and two-thirds of Shafin et al., 2026); "
                "if C4 ≤ C5, adaptivity has no value; if delegation precision is at "
                "chance, the claimed mechanism is absent. Every outcome is "
                "publishable."),
            "Risks and mitigations": (
                "Budget equalization may distort coordination-heavy arms (mitigated "
                "by three tiers and full budget–performance curves); model drift "
                "(version pinning, repetitions, paired analysis); contamination "
                "(temporal holdout); a weak controller (report the counterfactual "
                "oracle upper bound — itself a contribution)."),
        },
        "references": [
            "Agha, M., & Miqdad, A. (2026). An empirical comparison of multi-agent "
            "LLM architectural patterns for automated unit test generation "
            "[Master's thesis, Chalmers University of Technology and University of "
            "Gothenburg].",
            "Huang, J., Ye, W., Sun, W., Feng, Y., & Liu, Y. (2026). Cross-model "
            "collaboration for enhancing LLM-based code generation. ACM "
            "Transactions on Software Engineering and Methodology. "
            "https://doi.org/10.1145/3840382",
            "Li, Y. (2026). A multi-agent LLM framework for automated software "
            "testing. Transactions on Computing Science, 2(2). "
            "https://doi.org/10.63808/tcs.v2i2.447",
            "Orogat, A., Rostam, A., & Mansour, E. (2026). Understanding "
            "multi-agent LLM frameworks: A unified benchmark and experimental "
            "analysis. arXiv preprint arXiv:2602.03128. [preprint]",
            "Radeva, I., Noncheva, T., Doukovska, L., & Popchev, I. (2026). "
            "Comparing single-agent and multi-agent strategies in LLM-based "
            "title-abstract screening. Preprints.org. "
            "https://doi.org/10.20944/preprints202603.2107.v1 [preprint]",
            "Sanabria, D. (2026). OpenAI single-agent LLM architecture reduces "
            "computational overhead relative to multi-agent orchestration in a "
            "simulated Mars rover decision-support benchmark. Frontiers in "
            "Robotics and AI. https://doi.org/10.3389/frobt.2026.1877762",
            "Shafin, W. I., Rafi, M. N., Li, Z., & Chen, T.-H. (2026). An "
            "empirical study of Waterfall-style multi-agent workflows for "
            "class-level code generation. PROMISE '26: Proceedings of the 22nd "
            "International Conference on Predictive Models and Data Analytics in "
            "Software Engineering. https://doi.org/10.1145/3803846.3807461",
            "Yang, Y., Chai, H., & Zhang, W. (2025). AgentNet: Decentralized "
            "evolutionary coordination for LLM-based multi-agent systems. Qeios. "
            "https://doi.org/10.32388/ws0vim [preprint]",
            "Zeng, Z., Li, Y., Xie, R., Ye, W., & Zhang, S. (2025). Benchmarking "
            "and studying the LLM-based agent system in end-to-end software "
            "development. arXiv preprint arXiv:2511.04064. [preprint]",
        ],
    },
    # ------------------------------------------------------------------ T2
    {
        "heading": "Topic 2 — Independent Signals or Expensive Self-Refinement?",
        "proposed_title": ("Independent Signals or Expensive Self-Refinement? A "
                           "Budget-Matched Factorial Decomposition of Multi-Agent "
                           "Verification Gains in LLM-Based Code Review and Testing"),
        "problem": [
            "Multi-agent LLM systems for software verification — "
            "reviewer/critic/tester pipelines layered over a code generator — "
            "routinely report accuracy gains, but the literature systematically "
            "confounds three distinct causal candidates: (1) the role-play "
            "apparatus itself (persona labels plus multi-turn discussion), "
            "(2) genuinely independent signals (a different model, an independent "
            "context uncontaminated by the generator's reasoning, or executable "
            "evidence from tests and compilers), and (3) simple inference-budget "
            "inflation.",

            "The flagship multi-role vulnerability-detection study runs every role "
            "on instances of the same gpt-3.5-turbo model and concedes its gains "
            "cost a 484% token increase with no equal-budget single-agent control "
            "(Mao et al., 2024). The most controlled specialization ablation in the "
            "corpus finds a 17.5-point detection advantage for role decomposition "
            "but calls it “the only controlled evidence in this work,” runs "
            "entirely within one model at roughly three times the single agent's "
            "tokens, and states that a multi-model evaluation would be required to "
            "separate framework effects from model effects (Li, 2026). Decision "
            "theory sharpens the worry: without new exogenous signals, any "
            "delegated agent network is dominated by a centralized decision maker "
            "with the same information (Ao, Gao, & Simchi-Levi, 2026). Same-family "
            "self-auditing risks shared blind spots (Calboreanu, 2026), and "
            "multi-agent systems can “unexpectedly underperform strong, yet "
            "simpler, single-agent baselines” (Barrak, 2025).",

            "It is therefore still unknown whether multi-agent verification gains "
            "come from role labels and discussion or from independent signals, and "
            "whether same-model role-play is anything more than expensive "
            "self-refinement.",
        ],
        "evidence": [
            ("The theoretical core — a formal decision-theoretic result.",
             "We show that, without new exogenous signals, any delegated network is "
             "decision-theoretically dominated by a centralized Bayes decision "
             "maker with access to the same information.",
             "Ao, Gao, & Simchi-Levi (2026), p. 1 [preprint]"),
            ("Deployed architectures add roles without adding information.",
             "Planner, worker, critic, and reviewer modules are often built from "
             "the same model family, operate on overlapping retrieved context, and "
             "communicate through free-form language. In such cases, the added "
             "stages mainly transform and relay shared evidence rather than expand "
             "the information available for the terminal decision.",
             "Ao et al. (2026), p. 2 [preprint]"),
            ("The canonical multi-role verification study is fully confounded — "
             "role discussion, same model, no equal-token control.",
             "In terms of computation costs, due to the need for conversation "
             "between different roles, it requires a 484% increase in the number of "
             "tokens consumed.",
             "Mao, Li, Jin, Li, & Tei (2024), p. 2, IEEE QRS-C 2024"),
            ("The tightest ablation still triples the budget and stays within one "
             "model (9,348 vs. 3,120 average tokens).",
             "The single-agent condition detects fewer defects than the full "
             "framework under every criterion, and the gap of 17.5 points under the "
             "differential criterion is the only controlled evidence in this work "
             "that role decomposition contributes to detection.",
             "Li (2026), p. 17"),
            ("The isolation experiment is explicitly missing.",
             "All conditions reported in Table 6 use GLM-4-Flash, so the ablation "
             "isolates architectural contributions only within that model. … a "
             "multi-model evaluation would be required to separate framework "
             "effects from model effects.",
             "Li (2026), p. 20"),
            ("Same-model verification proves consistency, not correctness.",
             "The same LLM family (Claude) both authored AEGIS's prompt "
             "specifications and ran the audit; the inspector and the author could "
             "in principle share blind spots, and a clean round-9 pass demonstrates "
             "only internal consistency under that single family's vantage point, "
             "not ground-truth correctness.",
             "Calboreanu (2026), p. 17, Software (MDPI)"),
            ("The sign of the effect is genuinely uncertain.",
             "recent work has shown that multi-agent systems can unexpectedly "
             "underperform strong, yet simpler, single-agent baselines due to "
             "coordination overhead and these very cascading errors.",
             "Barrak (2025), p. 1, IEEE/ACM ASEW 2025"),
            ("Independent confirmation from the requirements domain.",
             "Yet, few studies experimentally isolate the marginal effect of model "
             "plurality from iterative workflow, or tie interaction patterns … to "
             "externally judged requirement value.",
             "Fan, Liu, Pan, Zhang, & Guo (2026), p. 4, Proc. ACM Softw. Eng. (FSE)"),
        ],
        "experiment": {
            "Design": (
                "A frozen-artifact, budget-matched factorial experiment. Stage 1: "
                "one fixed generator produces candidate solutions/patches, frozen "
                "and reused identically across all verification arms — arms differ "
                "only in how the artifact is verified. Stage 2: a 2×2×2×2 factorial "
                "manipulates four binary signal factors, with two anchored "
                "baselines: B1 generator-only (floor) and B2 same-model "
                "self-refinement at the same total token budget and round cap — the "
                "critical equal-budget comparator that Mao et al. (2024) and Li "
                "(2026) lack. Nineteen arms total (18 core plus a two-reviewer "
                "union arm testing signal complementarity per Calboreanu, 2026)."),
            "Conditions": [
                "F1 — reviewer model identity: same model as generator vs. a "
                "capability-matched different-family model",
                "F2 — context independence: reviewer sees the generator's full "
                "reasoning transcript vs. only the task spec and final artifact",
                "F3 — role framing: role-labeled persona with multi-turn discussion "
                "vs. a plain unpersonified check prompt",
                "F4 — executable evidence: sandboxed test/compiler output provided "
                "vs. none",
                "Baselines: B1 generator-only; B2 equal-budget same-model "
                "self-refinement",
            ],
            "Tasks and data": (
                "HumanEval+/MBPP+ (EvalPlus; 20% visible tests, hidden-suite "
                "scoring per Huang et al., 2026), QuixBugs with the two-run "
                "differential protocol (Li, 2026), a 50-task SWE-bench Lite subset, "
                "plus mutation-injected faults for ground-truth review "
                "precision/recall. Model versions pinned, five seeds, and a "
                "direction-swapped replication leg."),
            "Controls": (
                "Frozen artifacts, byte-identical prompts except the manipulated "
                "factor block, equal round caps and token ceilings, pinned sandbox "
                "— the four signal factors are the only free variables."),
            "Metrics": [
                "Hidden-suite pass@1; differential and strict defect detection",
                "Review precision/recall on injected faults",
                "Harm rate: initially-correct artifacts made incorrect by review",
                "Tokens, calls, latency, cost (verifying budget matching held)",
                "Generator–reviewer error-set overlap as a mediator",
                "Mixed-effects logistic regression on F1–F4 with Holm-corrected "
                "contrasts against B2",
            ],
            "Feasibility": (
                "Approximately 550 tasks × 19 arms × 5 seeds ≈ 500M tokens ≈ "
                "$500–$2,000 at small-model pricing; the frozen-artifact design "
                "amortizes generation across all arms; all harnesses are "
                "open-source; local-model replication is precedented (1,080 runs on "
                "quantized local models — Tomic, Alégroth, & Isaac, 2025). The "
                "cheapest and fastest of the three topics to execute."),
        },
        "viability": [
            ("The central manipulation is already proven executable.",
             "is this diversity necessary for CMCS's success, or does the "
             "improvement stem from the collaborative framework itself rather than "
             "model diversity? To answer this question, we conducted a controlled "
             "comparison between cross-model and same-model collaboration.",
             "Huang et al. (2026), p. 19, ACM TOSEM"),
            ("The executable-evidence factor has a demonstrated effect.",
             "Both Error Msgs and CoR thrive on the feedback from code compilers "
             "and achieve more than 1.2% improvements than Self-Refine, "
             "demonstrating that compilers can provide valuable signals to help "
             "LLMs better recognize the code bugs.",
             "Wang et al. (2024), p. 6, Findings of ACL 2024"),
            ("The statistical machinery transfers directly.",
             "To separate plurality from provider capability, we estimate a "
             "difference-in-differences (DiD) on session-level outcomes using the "
             "same-model controls (C1-MV, C3-SM), with BCa bootstrap CIs",
             "Fan et al. (2026), p. 12, Proc. ACM Softw. Eng. (FSE)"),
            ("Model-identity independence is a real, isolable signal.",
             "The two blind spots are vendor-specific and non-overlapping, which is "
             "itself the case for a multi-vendor panel: no single model is "
             "uniformly best, but the union of the panel detects every defect.",
             "Calboreanu (2026), p. 13, Software (MDPI)"),
            ("Dissertation-scale factorials can produce publishable negative "
             "results.",
             "We cannot accept our hypothesis H1, as one-model MALLMs tend to "
             "perform better on specific websites.",
             "Tomic, Alégroth, & Isaac (2025), p. 9, IEEE ICST 2025"),
        ],
        "assessment": {
            "Novelty": (
                "No corpus paper crosses the four signal factors, and none matches "
                "budgets. Li (2026) defers exactly this study; Huang et al. (2026) "
                "isolate one factor without budget-matched single-agent controls; "
                "Mao et al. (2024) have no control at all."),
            "Falsifiability": (
                "H1 (role framing adds nothing over budget-matched self-refinement "
                "once independent signals are controlled) dies if F3 shows a "
                "significant positive main effect against B2. H2 (gains concentrate "
                "in F1/F2/F4) dies on null independent-signal effects. Corpus "
                "precedent exists for both outcomes."),
            "Risks and mitigations": (
                "Small effects versus seed noise (pilot power analysis, equivalence "
                "bounds); tokenizer mismatch across families (match on normalized "
                "cost, report both); contamination (QuixBugs differential protocol "
                "plus mutation faults); capability mismatch masquerading as an "
                "identity effect (calibration within 3 pass@1 points plus a "
                "direction-swapped replication)."),
        },
        "references": [
            "Ao, R., Gao, S., & Simchi-Levi, D. (2026). On the reliability limits "
            "of LLM-based multi-agent planning. arXiv/SSRN preprint. "
            "https://doi.org/10.2139/ssrn.6490578 [preprint]",
            "Barrak, A. (2025). Traceability and accountability in "
            "role-specialized multi-agent LLM pipelines. 2025 40th IEEE/ACM "
            "International Conference on Automated Software Engineering Workshops "
            "(ASEW). https://doi.org/10.1109/asew67777.2025.00064",
            "Calboreanu, E. (2026). Iterative audit convergence in LLM-managed "
            "multi-agent systems: A case study in prompt-engineering quality "
            "assurance. Software, 5(2). https://doi.org/10.3390/software5020026",
            "Fan, G., Liu, D., Pan, L., Zhang, R., & Guo, Q. (2026). Multi-LLM "
            "persona generation for virtual focus groups in software engineering: "
            "A controlled, multi-domain study of emotional requirements "
            "elicitation. Proceedings of the ACM on Software Engineering, 3 (FSE). "
            "https://doi.org/10.1145/3808098",
            "Huang, J., Ye, W., Sun, W., Feng, Y., & Liu, Y. (2026). Cross-model "
            "collaboration for enhancing LLM-based code generation. ACM "
            "Transactions on Software Engineering and Methodology. "
            "https://doi.org/10.1145/3840382",
            "Li, Y. (2026). A multi-agent LLM framework for automated software "
            "testing. Transactions on Computing Science, 2(2). "
            "https://doi.org/10.63808/tcs.v2i2.447",
            "Mao, Z., Li, J., Jin, D., Li, M., & Tei, K. (2024). Multi-role "
            "consensus through LLMs discussions for vulnerability detection. 2024 "
            "IEEE International Conference on Software Quality, Reliability, and "
            "Security Companion (QRS-C). "
            "https://doi.org/10.1109/qrs-c63300.2024.00173",
            "Tomic, S., Alégroth, E., & Isaac, M. (2025). Evaluation of the "
            "choice of LLM in a multi-agent solution for GUI-test generation. 2025 "
            "IEEE Conference on Software Testing, Verification and Validation "
            "(ICST). https://doi.org/10.1109/icst62969.2025.10989038",
            "Wang, H., Liu, Z., Wang, S., Cui, G., Ding, N., Liu, Z., & Yu, G. "
            "(2024). INTERVENOR: Prompting the coding ability of large language "
            "models with the interactive chain of repair. Findings of the "
            "Association for Computational Linguistics: ACL 2024. "
            "https://doi.org/10.18653/v1/2024.findings-acl.124",
        ],
    },
    # ------------------------------------------------------------------ T3
    {
        "heading": "Topic 3 — Information-Preserving Handoff Contracts",
        "proposed_title": ("Information-Preserving Handoff Contracts for Multi-Agent "
                           "LLM Software Workflows: Isolating the Inter-Agent Artifact "
                           "Format as the Causal Variable in Requirement Preservation "
                           "and Error Propagation"),
        "problem": [
            "Role-specialized multi-agent LLM pipelines pass work between stages "
            "through ad hoc, mostly free-form prose handoffs, and a converging body "
            "of evidence shows that this relay step — not model capability — is "
            "where decision-relevant information dies. Controlled mechanism "
            "experiments show that inserting prose relay stages between an "
            "identical model and its answer collapses accuracy from 90.7% to 41.2% "
            "at two stages and 22.5% at five (Ao et al., 2026). In end-to-end "
            "software development, 55.8% of failures originate in planning-stage "
            "requirement omission and misinterpretation, with a specific handoff "
            "pathology: a flawed upstream design artifact is treated as "
            "authoritative and its errors propagate faithfully even when the "
            "original requirements remain in context (Zeng et al., 2025). In 100% "
            "of the 41 observed hierarchical failures in a controlled "
            "three-architecture comparison, correct worker reasoning never reached "
            "the worker that needed it (Agha & Miqdad, 2026). The problem is now a "
            "named production defect class — “Knowledge Transmission Bug” — in a "
            "taxonomy of 998 framework issue reports (Zhu et al., 2026), and the "
            "field's flagship review lists communication bottlenecks and "
            "information overload as unresolved scaling barriers (He, Treude, & "
            "Lo, 2025).",

            "What makes this dissertation-grade is that the design space is "
            "demonstrably non-monotonic — structured relays helped in some "
            "settings (Ao et al., 2026; a 36.22-point gain in Barrak, 2025) while "
            "rigid schemas actively harmed in others, converting correct reasoning "
            "into parse failures (Orogat et al., 2026) — and no study has isolated "
            "the handoff contract as the sole experimental variable on executable "
            "software tasks. Practitioners have no evidence-based guidance on what "
            "an inter-agent artifact must contain.",
        ],
        "evidence": [
            ("Relay steps destroy information even with identical models.",
             "Relative to the centralized baseline (A), the relay conditions (B2), "
             "(B), and (B5) add stages without adding new exogenous signals, and "
             "their performance is substantially lower. On gpt-4.1-mini, accuracy "
             "falls from 90.7% in (A) to 41.2% in (B2), 43.5% in (B), and 22.5% in "
             "(B5).",
             "Ao et al. (2026), p. 17 [preprint]"),
            ("Flawed upstream artifacts become authoritative at the handoff.",
             "Once the Dev Agent receives what it perceives as an authoritative "
             "implementation plan, it tends to prioritize this plan over direct "
             "engagement with the original requirement document, even though the "
             "latter is provided as context. If the design plan itself is flawed, "
             "such errors are faithfully propagated downstream, resulting in a "
             "final product that diverges from the actual requirements.",
             "Zeng et al. (2025), p. 8 [preprint]"),
            ("Requirement preservation is the binding constraint.",
             "The results in Table 5, clearly indicate that 'Task Planning' is the "
             "primary bottleneck in current agent systems, accounting for 55.8% of "
             "all issues.",
             "Zeng et al. (2025), p. 9 [preprint]"),
            ("The information bottleneck is total in hierarchical failures.",
             "No direct information flow between workers was observed in any of "
             "the 41 hierarchical failures in the complete dataset.",
             "Agha & Miqdad (2026), p. 72"),
            ("Errors cascade without provenance.",
             "When agents are chained together, an error introduced by an early "
             "agent can silently cascade, corrupting the entire workflow and "
             "leading to a final output that is incorrect for reasons that are "
             "difficult to diagnose.",
             "Barrak (2025), p. 1, IEEE/ACM ASEW 2025"),
            ("It is a named production defect class (998 issue reports).",
             "R11. Knowledge Transmission Bug (2.10%). It emerges in multi-agent "
             "settings when the framework fails to correctly propagate context, "
             "task specifications, or role-specific information across agents.",
             "Zhu et al. (2026), p. 3, FSE Companion '26"),
            ("The field's flagship review names it unresolved.",
             "Second, as the number of agents increases, so does the complexity of "
             "communication. Coordinating multiple agents can lead to "
             "communication bottlenecks and information overload.",
             "He, Treude, & Lo (2025), p. 19, ACM TOSEM"),
        ],
        "experiment": {
            "Design": (
                "A preregistered, randomized factorial experiment in which the only "
                "manipulated variable is the inter-agent handoff contract inside a "
                "fixed sequential pipeline (Planner→Coder→Tester primary; "
                "Developer→Tester replication). All conditions share one pinned "
                "model, byte-identical role prompts except the artifact-definition "
                "paragraph, identical per-stage step and token ceilings, and — "
                "critically — an identical token cap on the handoff artifact "
                "itself, so conditions differ in what the artifact contains and how "
                "it is organized, not how much can be said."),
            "Conditions": [
                "C0: budget-matched single agent (no handoff — the pipeline can "
                "lose outright)",
                "C1: free-form prose relay (status quo)",
                "C2: raw-transcript relay truncated to the same cap (curation vs. "
                "volume)",
                "C3: full structured contract — requirement checklist with IDs and "
                "verbatim requirement text (provenance anchors), design decisions "
                "with rationale, acceptance criteria with input/expected-output "
                "pairs, verification-performed log, and open-questions/uncertainty "
                "flags; schema-validated with one repair retry",
                "C4: C3 minus provenance (isolates the provenance field)",
                "C5: C3 minus uncertainty flags (isolates doubt signaling)",
                "C6: C3 format without schema validation (format vs. enforcement — "
                "resolves the confound in Barrak, 2025)",
                "Crossed factor: original requirements passed downstream alongside "
                "the artifact vs. replaced by it — the first direct test of Zeng et "
                "al.'s authoritative-plan mechanism",
                "Seeded-flaw sub-experiment: a controlled defect injected into the "
                "handoff artifact; the outcome is downstream detection vs. "
                "propagation",
            ],
            "Tasks and data": (
                "40–50 requirement-dense project-generation tasks per the "
                "E2EDevBench protocol (post-cutoff PyPI packages, atomic "
                "requirement checklists, migrated original test suites), with "
                "ClassEval and TestEval replication families. Local pytest "
                "sandbox."),
            "Controls": (
                "Pinned model and prompts, identical stage budgets, and an "
                "identical token cap on the handoff artifact across all "
                "conditions; the contract format is the sole free variable."),
            "Metrics": [
                "Per-requirement implementation rate (migrated tests, three-judge "
                "unanimous LLM verification, 10% human audit)",
                "Failure-mode distribution (Zeng et al. taxonomy)",
                "Handoff information-retention score: a blinded probe model "
                "reconstructs the requirement checklist from the artifact alone",
                "Seeded-flaw detection vs. propagation; parse-failure rate",
                "Full cost accounting; mixed-effects logistic regression with "
                "Holm–Bonferroni correction; preregistered power analysis "
                "(~50 tasks × 5 runs detects ~10-point differences)",
            ],
            "Feasibility": (
                "Entirely programmatic; public or reconstructible benchmarks; "
                "$3,000–$6,000 estimated API spend (corpus anchors: ~$7 per "
                "project worst case, $0.08–$0.31 per task typical); the harness is "
                "buildable by one student in a semester; runs are parallel and "
                "resumable."),
        },
        "viability": [
            ("The handoff format is causally manipulable with large, detectable "
             "effects.",
             "Condition (B post), which uses a more structured posterior-style "
             "message format, performs better than condition (B), which uses "
             "free-form prose, reaching 75.2% rather than 58.1% at three stages "
             "(on 50 questions; see Figure 3).",
             "Ao et al. (2026), p. 17 [preprint]"),
            ("The intervention class works at published scale; that study's "
             "confounds (prompts changed with protocol, MCQ-only, no repeated "
             "runs) define exactly what this design isolates.",
             "On the PythonIO benchmark, the accuracy gains were even more "
             "substantial for several configurations. The BBB configuration's "
             "accuracy increased by 36.22 percentage points, from 61.42% to "
             "97.64%.",
             "Barrak (2025), p. 5, IEEE/ACM ASEW 2025"),
            ("Agents reliably produce and consume such contracts already.",
             "In 87 of 90 runs across the entire dataset, the planner agent "
             "produced a structured test specification that included explicit "
             "input values, expected outputs, and the specific code paths each "
             "test case targeted.",
             "Agha & Miqdad (2026), p. 63"),
            ("The controlled-comparison template exists.",
             "Using this framework, we conduct a controlled empirical study on "
             "three representative agent architectures implemented upon a unified "
             "foundation to isolate the impact of workflow design.",
             "Zeng et al. (2025), p. 1 [preprint]"),
            ("The outcome is genuinely uncertain in both directions "
             "(falsifiable).",
             "Table 9 shows that schema-constrained planning reduces accuracy "
             "across datasets and models, while free-form planning preserves or "
             "improves accuracy in most settings.",
             "Orogat et al. (2026), p. 8 [preprint]"),
        ],
        "assessment": {
            "Novelty": (
                "No corpus paper isolates the handoff contract as the sole "
                "variable on executable software tasks, ablates contract fields "
                "(provenance, uncertainty, validation), token-matches the artifact "
                "across formats, or measures information retention and seeded-flaw "
                "propagation as outcomes. Zeng et al.'s authoritative-plan "
                "mechanism is an untested hypothesis this design tests directly."),
            "Falsifiability": (
                "H1 fails if token-matched structured contracts do not beat prose "
                "(a real possibility — Orogat et al. document schema-induced "
                "collapse); H2 fails if the provenance ablation shows nothing; C0 "
                "lets the pipeline lose outright to a single agent."),
            "Risks and mitigations": (
                "LLM-judge validity (migrated original tests primary, unanimous "
                "judges, human audit); contamination (post-cutoff PyPI selection); "
                "model drift (pinned snapshots, second-family replication); "
                "schema-encodes-task-knowledge circularity (token cap plus the "
                "raw-transcript control C2)."),
        },
        "references": [
            "Agha, M., & Miqdad, A. (2026). An empirical comparison of multi-agent "
            "LLM architectural patterns for automated unit test generation "
            "[Master's thesis, Chalmers University of Technology and University of "
            "Gothenburg].",
            "Ao, R., Gao, S., & Simchi-Levi, D. (2026). On the reliability limits "
            "of LLM-based multi-agent planning. arXiv/SSRN preprint. "
            "https://doi.org/10.2139/ssrn.6490578 [preprint]",
            "Barrak, A. (2025). Traceability and accountability in "
            "role-specialized multi-agent LLM pipelines. 2025 40th IEEE/ACM "
            "International Conference on Automated Software Engineering Workshops "
            "(ASEW). https://doi.org/10.1109/asew67777.2025.00064",
            "He, J., Treude, C., & Lo, D. (2025). LLM-based multi-agent systems "
            "for software engineering: Literature review, vision, and the road "
            "ahead. ACM Transactions on Software Engineering and Methodology. "
            "https://doi.org/10.1145/3712003",
            "Orogat, A., Rostam, A., & Mansour, E. (2026). Understanding "
            "multi-agent LLM frameworks: A unified benchmark and experimental "
            "analysis. arXiv preprint arXiv:2602.03128. [preprint]",
            "Zeng, Z., Li, Y., Xie, R., Ye, W., & Zhang, S. (2025). Benchmarking "
            "and studying the LLM-based agent system in end-to-end software "
            "development. arXiv preprint arXiv:2511.04064. [preprint]",
            "Zhu, X., Wu, J., Zhang, X., Li, T., Mu, Y., Zhai, J., Shen, C., "
            "Fang, C., & Liu, Y. (2026). Bugs in modern LLM agent frameworks: An "
            "empirical study. FSE Companion '26: Proceedings of the 34th ACM "
            "International Conference on the Foundations of Software Engineering. "
            "https://doi.org/10.1145/3803437.3805536",
        ],
    },
]

CLOSING_HEADING = "Closing Notes"
CLOSING_PARAS = [
    ("Runner-up topic. A fourth candidate — semantic fault containment under "
     "matched architectures (whether specialist coordination creates net "
     "resilience or a larger cascade surface when upstream messages are fluent "
     "but semantically wrong) — ranked closely behind the three presented here. "
     "It was set aside because the fault-injection harness pattern is already "
     "established in the literature and its cleanest formulation overlaps Topic "
     "3's seeded-flaw sub-experiment. If pursued, it should be extended with an "
     "explicit adversarial arm: security is the corpus's largest verified "
     "absence — none of the major agent benchmarks reviewed scores security, and "
     "hierarchical coordination uniquely produced delegation-leak outputs in the "
     "corpus's controlled comparison."),
    ("Cross-cutting methodology. Whichever topic is selected, two corpus-wide "
     "findings shape the experimental design. First, evaluation validity: "
     "headline success metrics in this literature inflate substantially under "
     "stricter criteria (detection rates collapsing as criteria tighten; a "
     "meaningful share of unit-test-passing solutions failing property-based "
     "tests), so the proposed experiments use tiered success criteria, migrated "
     "or property-based tests, and hidden suites. Second, run-to-run variance: "
     "most primary studies run one to five trials without intervals, while the "
     "corpus shows architecture shifts the outcome distribution, not just its "
     "mean — so all proposed designs preregister repeated runs and paired "
     "task-level inference."),
    ("Provenance. The complete audit trail behind this document — search logs, "
     "screening decisions with rationales, citation-snowball edges, the "
     "per-paper evidence ledger, the saturation report, and the extracted "
     "full-text corpus against which every quotation above was verified — is "
     "preserved in the project repository and available on request."),
]

# --------------------------------------------------------------------------
# Document construction
# --------------------------------------------------------------------------

doc = Document()

# Page geometry
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_font(style, size, bold=False, italic=False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")


set_font(doc.styles["Normal"], 12)
doc.styles["Normal"].paragraph_format.line_spacing = 1.0

for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
    set_font(doc.styles[name], size, bold=True, italic=(name == "Heading 3"))
    doc.styles[name].paragraph_format.space_before = Pt(18 if name == "Heading 1" else 12)
    doc.styles[name].paragraph_format.space_after = Pt(6)
    doc.styles[name].paragraph_format.keep_with_next = True


def body(text, first_indent=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 2.0
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(0)
    return p


def claim(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{num}. {text}")
    run.bold = True
    return p


def block_quote(text):
    p = doc.add_paragraph(f"“{text}”")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    return p


def attribution(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"— {text}")
    run.italic = True
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    set_font_run = p.runs[0] if p.runs else None
    return p


def labeled_para(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}. ")
    run.bold = True
    p.add_run(text)
    return p


def reference_entry(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(12)  # NSU: single-spaced, double space between
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
    return p


def add_page_number_footer():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    run.append(rpr)
    fld.append(run)
    p._p.append(fld)


# --- Title page ------------------------------------------------------------
for _ in range(6):
    doc.add_paragraph()
for line in TITLE.split("\n"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(line)
    run.bold = True
    run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
p.add_run(SUBTITLE).italic = True

for _ in range(3):
    doc.add_paragraph()
for line in [AUTHOR] + PROGRAM.split("\n") + [DATE]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0

doc.add_page_break()

# --- Opening section --------------------------------------------------------
doc.add_heading("Topic Area and Purpose of This Distillation", level=1)
for para in PURPOSE_PARAS:
    body(para)

# --- Topics ------------------------------------------------------------------
for topic in TOPICS:
    doc.add_page_break()
    doc.add_heading(topic["heading"], level=1)
    labeled_para("Proposed dissertation title", topic["proposed_title"])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    doc.add_heading("Problem Statement", level=2)
    for para in topic["problem"]:
        body(para)

    doc.add_heading("Evidence That This Is a Problem", level=2)
    for i, (claim_text, quote_text, attr) in enumerate(topic["evidence"], 1):
        claim(i, claim_text)
        block_quote(quote_text)
        attribution(attr)

    doc.add_heading("Proposed Experiment", level=2)
    exp = topic["experiment"]
    labeled_para("Design", exp["Design"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Conditions.")
    run.bold = True
    for c in exp["Conditions"]:
        bullet(c)
    labeled_para("Tasks and data", exp["Tasks and data"])
    labeled_para("Controls", exp["Controls"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Outcome metrics.")
    run.bold = True
    for m in exp["Metrics"]:
        bullet(m)
    labeled_para("Feasibility", exp["Feasibility"])

    doc.add_heading("Why the Proposed Approach Is Viable", level=2)
    for i, (claim_text, quote_text, attr) in enumerate(topic["viability"], 1):
        claim(i, claim_text)
        block_quote(quote_text)
        attribution(attr)

    doc.add_heading("Assessment", level=2)
    for label, text in topic["assessment"].items():
        labeled_para(label, text)

    doc.add_heading("References for This Topic", level=2)
    for ref in topic["references"]:
        reference_entry(ref)

# --- Closing -----------------------------------------------------------------
doc.add_page_break()
doc.add_heading(CLOSING_HEADING, level=1)
for para in CLOSING_PARAS:
    label, rest = para.split(". ", 1)
    labeled_para(label, rest)

add_page_number_footer()
doc.save(OUT)
print(f"Wrote {OUT}")
