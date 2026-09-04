# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx"]
# ///
"""Render a run's thesis.md into the NSU-format advisor-review Word document.

Consumes the markdown subset documented in
.github/skills/publications-search/references/THESIS-TEMPLATE.md: a leading H1
title, a bold **Author:**/**Date:**/**Program:**/**Purpose:** metadata
paragraph feeding the title page, "---" rules that close the current topic
scope, "# Topic ..." sections (each starts on a new page), and a trailing
"## References" section rendered with hanging indents. Formatting follows NSU
conventions: Times New Roman, double-spaced body with first-line indents, bold
numbered claims over indented block quotes with em-dash italic attributions,
single-spaced references, and a PAGE-field footer.
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

EXIT_SUCCESS = 0
EXIT_ERROR = 2

logger = logging.getLogger(__name__)

EM_DASH = "—"

RULE_RE = re.compile(r"^-{3,}\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^- +(.*)$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
META_KEY_RE = re.compile(r"\*\*([A-Za-z][A-Za-z ]*):\*\*\s*")

# Inline subset: **bold**, *italic*, `code` -> plain text, [text](url) -> text.
_INLINE_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<italic>[^*]+)\*"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<label>[^\]]+)\]\((?P<url>[^)]+)\)"
)

# Known topic subsection headings are renamed to their formal document titles;
# anything else renders under its own heading so future runs can add sections.
TOPIC_SECTION_TITLES = {
    "the problem": "Problem Statement",
    "evidence that this is a problem": "Evidence That This Is a Problem",
    "the experiment": "Proposed Experiment",
    "why this approach is viable": "Why the Proposed Approach Is Viable",
}


# --------------------------------------------------------------------------
# thesis.md parsing
# --------------------------------------------------------------------------


@dataclass
class Claim:
    number: int
    text: str
    quote: str | None = None
    attribution: str | None = None


@dataclass
class Block:
    kind: str  # paragraph | bullets | claims | quote
    text: str = ""
    items: list = field(default_factory=list)
    attribution: str | None = None


@dataclass
class Section:
    heading: str | None
    blocks: list[Block] = field(default_factory=list)


@dataclass
class Topic:
    heading: str
    blocks: list[Block] = field(default_factory=list)
    sections: list[Section] = field(default_factory=list)


@dataclass
class Thesis:
    title: str
    meta: dict[str, str]
    units: list[Topic | Section]


def _split_quote_lines(raw: list[str]) -> tuple[str | None, str | None]:
    """Blockquote lines are quote text until the first em-dash line, which
    starts the attribution; attributions may wrap onto further "> " lines."""
    quote: list[str] = []
    attrib: list[str] = []
    for text in raw:
        if attrib:
            attrib.append(text)
        elif text.startswith(EM_DASH):
            attrib.append(text.lstrip(EM_DASH).strip())
        else:
            quote.append(text)
    quote_text = " ".join(quote).strip().strip('"“”') or None
    return quote_text, (" ".join(attrib).strip() or None)


def tokenize(lines: list[str]) -> list[tuple[str, object]]:
    events: list[tuple[str, object]] = []
    i, n = 0, len(lines)

    def skip_blanks(j: int) -> int:
        while j < n and not lines[j].strip():
            j += 1
        return j

    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if RULE_RE.match(stripped):
            events.append(("rule", None))
            i += 1
            continue
        heading = HEADING_RE.match(lines[i])
        if heading:
            events.append((f"h{len(heading.group(1))}", heading.group(2).strip()))
            i += 1
            continue
        if BULLET_RE.match(lines[i]):
            items: list[str] = []
            while i < n:
                m = BULLET_RE.match(lines[i])
                if not m:
                    j = skip_blanks(i)
                    if j < n and BULLET_RE.match(lines[j]):
                        i = j
                        continue
                    break
                item = [m.group(1).strip()]
                i += 1
                # Continuation lines are indented and carry no bullet marker.
                while (
                    i < n
                    and lines[i].strip()
                    and lines[i].startswith(" ")
                    and not BULLET_RE.match(lines[i])
                ):
                    item.append(lines[i].strip())
                    i += 1
                items.append(" ".join(item))
            events.append(("bullets", items))
            continue
        if NUMBERED_RE.match(lines[i]):
            claims: list[Claim] = []
            while i < n:
                m = NUMBERED_RE.match(lines[i])
                if not m:
                    j = skip_blanks(i)
                    if j < n and NUMBERED_RE.match(lines[j]):
                        i = j
                        continue
                    break
                text_lines = [m.group(2).strip()]
                quote_lines: list[str] = []
                i += 1
                while i < n and lines[i].strip() and not NUMBERED_RE.match(lines[i]):
                    quoted = QUOTE_RE.match(lines[i])
                    if quoted and lines[i].lstrip().startswith(">"):
                        quote_lines.append(quoted.group(1).strip())
                    else:
                        text_lines.append(lines[i].strip())
                    i += 1
                quote, attribution = _split_quote_lines(quote_lines)
                claims.append(Claim(int(m.group(1)), " ".join(text_lines), quote, attribution))
            events.append(("claims", claims))
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quoted = QUOTE_RE.match(lines[i])
                quote_lines.append(quoted.group(1).strip() if quoted else "")
                i += 1
            events.append(("quote", _split_quote_lines(quote_lines)))
            continue
        para = [stripped]
        i += 1
        while i < n:
            line = lines[i]
            text = line.strip()
            if (
                not text
                or RULE_RE.match(text)
                or HEADING_RE.match(line)
                or BULLET_RE.match(line)
                or NUMBERED_RE.match(line)
                or text.startswith(">")
            ):
                break
            para.append(text)
            i += 1
        events.append(("paragraph", " ".join(para)))
    return events


def parse_meta(text: str) -> dict[str, str]:
    parts = META_KEY_RE.split(text)
    meta: dict[str, str] = {}
    for key, value in zip(parts[1::2], parts[2::2], strict=False):
        meta[key.strip().lower()] = value.strip().strip("·").strip()
    return meta


def parse_thesis(text: str) -> Thesis:
    events = tokenize(text.splitlines())
    title: str | None = None
    meta: dict[str, str] = {}
    units: list[Topic | Section] = []
    topic: Topic | None = None
    section: Section | None = None

    def target() -> list[Block]:
        if section is not None:
            return section.blocks
        if topic is not None:
            return topic.blocks
        # Content with no heading in scope (e.g. a closing note after the
        # final rule) collects into an anonymous trailing section.
        if not units or not isinstance(units[-1], Section) or units[-1].heading is not None:
            units.append(Section(None))
        return units[-1].blocks

    for kind, payload in events:
        if kind == "h1":
            if title is None:
                title = str(payload)
                continue
            topic = Topic(str(payload))
            units.append(topic)
            section = None
            continue
        if kind == "rule":
            # A rule closes the current topic: following H2s are top-level.
            topic = None
            section = None
            continue
        if kind.startswith("h"):
            section = Section(str(payload))
            if topic is not None:
                topic.sections.append(section)
            else:
                units.append(section)
            continue
        if (
            kind == "paragraph"
            and title is not None
            and not meta
            and not units
            and topic is None
            and META_KEY_RE.match(str(payload))
        ):
            meta = parse_meta(str(payload))
            continue
        if kind == "paragraph":
            target().append(Block("paragraph", text=str(payload)))
        elif kind in ("bullets", "claims"):
            target().append(Block(kind, items=list(payload)))
        elif kind == "quote":
            quote, attribution = payload
            target().append(Block("quote", text=quote or "", attribution=attribution))
    if title is None:
        raise ValueError("must start with an H1 title")
    return Thesis(title, meta, units)


def parse_inline(
    text: str, bold: bool = False, italic: bool = False
) -> list[tuple[str, bool, bool]]:
    runs: list[tuple[str, bool, bool]] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos : m.start()], bold, italic))
        if m.group("bold") is not None:
            runs.extend(parse_inline(m.group("bold"), True, italic))
        elif m.group("italic") is not None:
            runs.extend(parse_inline(m.group("italic"), bold, True))
        elif m.group("code") is not None:
            runs.append((m.group("code"), bold, italic))
        else:
            runs.append((m.group("label"), bold, italic))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], bold, italic))
    return [run for run in runs if run[0]]


def plain(text: str) -> str:
    return "".join(segment for segment, _, _ in parse_inline(text))


# --------------------------------------------------------------------------
# docx rendering (NSU formatting rules)
# --------------------------------------------------------------------------


def add_runs(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    for segment, seg_bold, seg_italic in parse_inline(text, bold, italic):
        run = paragraph.add_run(segment)
        run.bold = seg_bold
        run.italic = seg_italic


def set_font(style, size: int, bold: bool = False, italic: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")


def setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    set_font(doc.styles["Normal"], 12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.0
    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        set_font(doc.styles[name], size, bold=True, italic=(name == "Heading 3"))
        fmt = doc.styles[name].paragraph_format
        fmt.space_before = Pt(18 if name == "Heading 1" else 12)
        fmt.space_after = Pt(6)
        fmt.keep_with_next = True


def body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    # Bold-led paragraphs are labels (**Design.** / **Proposed title:**);
    # NSU labeled paragraphs take no first-line indent.
    if not text.startswith("**"):
        p.paragraph_format.first_line_indent = Inches(0.5)
    add_runs(p, text)
    return p


def claim(doc: Document, number: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    add_runs(p, f"{number}. {text}", bold=True)
    return p


def block_quote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_runs(p, f"“{text}”")
    return p


def attribution(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(10)
    p.add_run(f"{EM_DASH} ").italic = True
    add_runs(p, text, italic=True)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, text)
    return p


def reference_entry(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(12)  # NSU: single-spaced, double space between
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
    add_runs(p, text)
    return p


def add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    run.append(rpr)
    fld.append(run)
    p._p.append(fld)


def render_title_page(doc: Document, thesis: Thesis, args: argparse.Namespace) -> None:
    author = args.author or thesis.meta.get("author", "")
    program = args.program or thesis.meta.get("program", "")
    date = args.date or thesis.meta.get("date", "")
    purpose = thesis.meta.get("purpose", "")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(plain(thesis.title))
    run.bold = True
    run.font.size = Pt(16)
    if args.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        p.add_run(args.subtitle).italic = True
    for _ in range(3):
        doc.add_paragraph()
    # Program renders one comma-separated component per line (NSU title page).
    for line in [author, *program.split(", "), date]:
        if not line:
            continue
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
    if purpose:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        add_runs(p, purpose, italic=True)
    doc.add_page_break()


def section_mode(heading: str | None) -> str:
    key = (heading or "").lower()
    if key == "assessment":
        return "assessment"
    if key.startswith("references"):
        return "references"
    return "generic"


def render_blocks(doc: Document, blocks: list[Block], mode: str = "generic") -> None:
    for block in blocks:
        if block.kind == "paragraph":
            body_paragraph(doc, block.text)
        elif block.kind == "bullets":
            for item in block.items:
                if mode == "references":
                    reference_entry(doc, item)
                elif mode == "assessment":
                    body_paragraph(doc, item)  # labeled paragraphs, not bullets
                else:
                    bullet(doc, item)
        elif block.kind == "claims":
            for item in block.items:
                if item.quote or item.attribution:
                    claim(doc, item.number, item.text)
                    if item.quote:
                        block_quote(doc, item.quote)
                    if item.attribution:
                        attribution(doc, item.attribution)
                else:
                    body_paragraph(doc, f"{item.number}. {item.text}")
        elif block.kind == "quote":
            if block.text:
                block_quote(doc, block.text)
            if block.attribution:
                attribution(doc, block.attribution)


def render_topic(doc: Document, topic: Topic) -> None:
    doc.add_heading(plain(topic.heading), level=1)
    render_blocks(doc, topic.blocks)
    for section in topic.sections:
        if section.heading:
            title = TOPIC_SECTION_TITLES.get(section.heading.lower(), section.heading)
            doc.add_heading(plain(title), level=2)
        render_blocks(doc, section.blocks, section_mode(section.heading))


def render_document(thesis: Thesis, args: argparse.Namespace) -> Document:
    doc = Document()
    setup_document(doc)
    render_title_page(doc, thesis, args)
    for index, unit in enumerate(thesis.units):
        if isinstance(unit, Topic):
            if index > 0:
                doc.add_page_break()
            render_topic(doc, unit)
        else:
            if index > 0 and unit.heading:
                doc.add_page_break()
            if unit.heading:
                doc.add_heading(plain(unit.heading), level=1)
            render_blocks(doc, unit.blocks, section_mode(unit.heading))
    add_page_number_footer(doc)
    return doc


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------


def create_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("run_dir", type=Path, help="Run folder containing thesis.md.")
    parser.add_argument(
        "--thesis",
        default="thesis.md",
        help="Thesis markdown file, relative to run_dir unless absolute.",
    )
    parser.add_argument(
        "--out",
        default="Problem-Statement-Literature-Review.docx",
        help="Output .docx, written into run_dir unless absolute.",
    )
    parser.add_argument("--author", default=None, help="Override the **Author:** metadata.")
    parser.add_argument("--program", default=None, help="Override the **Program:** metadata.")
    parser.add_argument("--date", default=None, help="Override the **Date:** metadata.")
    parser.add_argument(
        "--subtitle",
        default="Submitted for Advisor Review",
        help="Title-page subtitle line (empty string omits it).",
    )
    parser.add_argument("-v", "--verbose", action="store_true")
    return parser


def main() -> int:
    args = create_parser().parse_args()
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s: %(message)s",
    )
    thesis_path = Path(args.thesis)
    if not thesis_path.is_absolute():
        thesis_path = args.run_dir / thesis_path
    if not thesis_path.is_file():
        logger.error("No thesis markdown at %s.", thesis_path)
        return EXIT_ERROR
    try:
        thesis = parse_thesis(thesis_path.read_text(encoding="utf-8"))
    except ValueError as exc:
        logger.error("%s: %s", thesis_path, exc)
        return EXIT_ERROR

    doc = render_document(thesis, args)
    out_path = Path(args.out)
    if not out_path.is_absolute():
        out_path = args.run_dir / out_path
    doc.save(str(out_path))
    topics = sum(1 for unit in thesis.units if isinstance(unit, Topic))
    logger.info(
        "%d topics, %d top-level sections, %d paragraphs -> %s",
        topics,
        len(thesis.units) - topics,
        len(doc.paragraphs),
        out_path,
    )
    return EXIT_SUCCESS


if __name__ == "__main__":
    sys.exit(main())
